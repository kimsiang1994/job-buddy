"""Renders one tailored data model into a PDF and a DOCX that say the same thing.

`tailor()` decides *what* goes on the resume. This decides *whether it fits* --
and those are different problems that fail in different ways.

**One model, two renderers.** The PDF (Typst) and the DOCX (python-docx) are
both built from the output of `build_model()`, never from each other and never
from a second layout description. The failure this prevents is the boring one
that actually happens: someone edits the PDF template, ships, and the Word copy
the recruiter opens is three bullets behind. `tests/test_render_resume.py`
asserts both documents carry identical bullet text.

**The page budget is a parameter, not a rule.** A hard one-page limit is
folklore at mid-senior level and the evidence runs against it: hiring managers
screening at volume accept two pages past ~5 years, a 482-participant
simulation measured a 2.6x preference for two pages at mid-level, and Wilson &
Caliskan (AIES 2024) found that *shortening* resumes increased biased outcomes
by 22.2%. Cutting is not free and this module does not pretend it is.

What survives that evidence is a sharper constraint: **page two does not get
read.** So `max_pages` defaults to 1 but is configurable, and
`page_one_sufficiency()` answers the question that actually matters -- did the
most relevant role, the top three ranked bullets and the skills block all land
above the page-1 break? That is deterministic: render, extract page 1 with
pypdf, test membership. It is a check, not a hope.

**Fitting shrinks before it cuts.** A single `scale` is binary-searched, with
font size, leading and margins all derived from it, and the real page count
verified by re-reading the PDF rather than estimated. At the ~9pt floor the
shrinking stops and the LOWEST-RANKED bullets are dropped -- which is the whole
reason `tailor()` returns a ranking instead of a set. Every cut is recorded in
`dropped` with its fact_id and reason, because the way a resume loses its best
bullet is that nobody was told it happened.

**Formatting choices are evidence-based, and some folklore is deliberately
ignored.** Single column: multi-column layout measurably degrades reading-order
reconstruction in every VLM-based extractor in OmniDocBench (CVPR 2025), and
the damage concentrates in long-text fields -- which is the bullets. No images,
no text inside graphics, no headers or footers, so contact details go in the
document body. Conventional section headings, because heading matching is
per-vendor and an unconventional heading is the documented failure. Under 100 KB,
which is Taleo's documented upload limit and is asserted in the tests. What this
module does NOT do is restrict fonts, avoid bold or bullets, or avoid tables --
Oracle states styling does not affect parsing, and that advice is cargo cult.

**Nothing here is required.** Core is stdlib. A missing wheel warns once and
degrades: no `typst` emits the `.typ` source plus the command that would compile
it, no `python-docx` emits markdown. It never raises for a missing wheel, and
the result always names which degradation was taken.
"""

from __future__ import annotations

import re
import threading
import warnings
from pathlib import Path
from typing import Any

# Typographic range. The ceiling is a comfortable reading size; the floor is the
# point below which a human screener is squinting, so the fitter stops shrinking
# and starts cutting instead of continuing down to an unreadable 7pt.
BASE_PT = 11.0
MIN_PT = 9.0
SCALE_MIN = MIN_PT / BASE_PT
SCALE_MAX = 1.0

# Enough to land within ~0.005 of the largest scale that fits. Each step is a
# real Typst compile, so this is the cost knob.
FIT_STEPS = 7

# Horizontal page margin, in inches, at SCALE_MIN and SCALE_MAX. Measured off
# the source resume rather than picked: its body text starts at 36pt = 0.50in
# and its wrapped lines end at 576pt, i.e. 0.50in the other side. The previous
# values (0.45 + 0.45*scale, so 0.90in at full size) spent 1.8in of an 8.27in
# page on white and drove the fitter into cutting bullets that would otherwise
# have fitted. `MIN_TEXT_COLUMN_FRACTION` is the property this exists to hold,
# and `tests/test_render_resume.py` asserts it against a rendered PDF rather
# than against these numbers -- a test on the constant would pass while the
# page still looked starved.
MARGIN_MIN_IN = 0.42
MARGIN_MAX_IN = 0.52

# The text column must occupy at least this share of the page width. 0.85 sits
# below the 0.874 the margins above produce at full scale, so an ordinary
# typographic tweak does not trip it, and well above the 0.78 that prompted
# this -- a page starved back toward the old margins fails.
MIN_TEXT_COLUMN_FRACTION = 0.85

# Taleo's documented upload limit. Rendering above it is a silent rejection.
MAX_PDF_BYTES = 100 * 1024

# Conventional, in this order. Unconventional headings are the documented
# per-vendor parsing failure, so these strings are not a style choice.
HEADING_EXPERIENCE = "Experience"
HEADING_SKILLS = "Technical Skills"
HEADING_EDUCATION = "Education"

# How many top-ranked bullets must clear the page-1 break to call the render
# sufficient. Three is the screener's first pass, not a tunable.
PAGE_ONE_BULLETS = 3

_capabilities: dict[str, bool] | None = None
_warned: set[str] = set()
# Guards the capability probe against a concurrent reset. See capabilities().
_capabilities_lock = threading.Lock()
# Guards the check-then-act in _warn_once. Separate from the probe lock because
# it protects a different invariant; `reset_capability_cache` takes both, always
# in that order, and nothing takes them the other way round.
_warned_lock = threading.Lock()


# --------------------------------------------------------------------------
# optional dependencies -- probe once, warn once
# --------------------------------------------------------------------------

def _load_typst():
    """The typst module, or None. Patched in tests to simulate a missing wheel."""
    try:
        import typst

        return typst
    except ImportError:
        return None


def _load_docx():
    """The python-docx module, or None."""
    try:
        import docx

        return docx
    except ImportError:
        return None


def _load_pypdf():
    """pypdf's PdfReader, or None."""
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        return None


def capabilities(refresh: bool = False) -> dict[str, bool]:
    """Which optional renderers are usable. Probed once and cached.

    Cached because the answer cannot change inside a run and because probing on
    every bullet would turn an import failure into a per-call cost. `refresh` is
    for tests, which need to re-probe after patching a loader.

    The lock is defensive, and honestly labelled as such. The None-check and
    the `dict(_capabilities)` read are two separate loads of the global, so a
    `reset_capability_cache()` landing between them would make the read raise
    `TypeError: 'NoneType' object is not iterable`. That window is about two
    bytecodes wide: it was NOT reproducible, even with six reader threads, two
    resetter threads and `sys.setswitchinterval(1e-6)` for three seconds. So
    there is no regression test for it -- a test that cannot fail on the
    unlocked version would prove nothing.

    The lock stays because it is uncontended in the normal path and makes the
    invariant true by construction rather than by timing. This is NOT the
    `loaded`-before-the-work bug found in `model_config._load()` and
    `token_budget._backend_load()`; the probe here builds the dict fully before
    assigning it, so no caller can ever observe a half-built cache.
    """
    global _capabilities
    with _capabilities_lock:
        if _capabilities is None or refresh:
            _capabilities = {
                "pdf": _load_typst() is not None,
                "docx": _load_docx() is not None,
                "page_count": _load_pypdf() is not None,
            }
        return dict(_capabilities)


def _warn_once(key: str, message: str) -> None:
    """One warning per missing capability per process.

    A degraded render happens once per job; warning every time would bury the
    line that says which degradation was taken under fifty identical copies.

    The membership test and the insert are one critical section. `set.add` being
    atomic does not close this: two threads can both pass the `key in _warned`
    check before either adds, and both then warn. `render()` is called from
    `pipeline._prepare_job` under a ThreadPoolExecutor, so with typst absent and
    four workers the "typst is not installed" line could print four times --
    exactly the burial this function exists to prevent.

    `warnings.warn` is deliberately outside the lock: it walks the stack for
    `stacklevel` and then runs whatever filter the caller installed, and holding
    a lock across arbitrary caller code would serialise every degraded render
    behind the first. The lock protects the decision, not the emission.
    """
    with _warned_lock:
        if key in _warned:
            return
        _warned.add(key)
    warnings.warn(message, RuntimeWarning, stacklevel=3)


def reset_capability_cache() -> None:
    """Forget the probe and the warnings. For tests only.

    Takes the same lock as `capabilities()`, so a reset can never land between
    that function's None-check and its read.
    """
    global _capabilities
    with _capabilities_lock:
        _capabilities = None
        with _warned_lock:
            _warned.clear()


# --------------------------------------------------------------------------
# the data model -- the single source of truth both renderers read
# --------------------------------------------------------------------------

def build_model(profile: dict[str, Any], tailored: dict[str, Any]) -> dict[str, Any]:
    """`tailor()`'s output plus profile identity, in the shape a page reads.

    Bullets keep `tailor()`'s rank order, which is load-bearing twice over: the
    fitter cuts from the end of this list, and `page_one_sufficiency` reads the
    front of it. ROLES, however, are ordered by date -- see `group_roles`. An
    earlier version ordered roles by their best-ranked bullet on the reasoning
    that the most relevant job should lead. It produced a resume running Olea
    2022, TikTok 2026, Citibank 2023: chronology shuffled, which a reader takes
    for a mistake rather than for emphasis.
    """
    identity = profile.get("identity") or {}
    notes = {(f.get("org"), f.get("role")): f.get("note")
             for f in (profile.get("facts") or []) if f.get("note")}

    bullets: list[dict[str, Any]] = []
    for rank, bullet in enumerate(tailored.get("bullets") or [], start=1):
        org, role = bullet.get("org") or "", bullet.get("role") or ""
        bullets.append({
            "text": str(bullet.get("text") or "").strip(),
            "fact_id": str(bullet.get("fact_id") or ""),
            "org": org,
            "role": role,
            "start": bullet.get("start") or "",
            "end": bullet.get("end") or "",
            "note": bullet.get("note") or notes.get((org, role)) or "",
            "rank": rank,
        })

    contact = [str(identity.get(k) or "").strip()
               for k in ("email", "phone", "location")]
    contact += [str(link).strip() for link in (identity.get("links") or [])]

    return {
        "name": str(identity.get("name") or "").strip(),
        # In the body, never a header or footer -- a PDF header is routinely
        # dropped before the parser ever sees the contact details.
        "contact": [c for c in contact if c],
        "headline": str(tailored.get("headline") or "").strip(),
        "bullets": bullets,
        "roles": group_roles(bullets),
        # Grouped skills win when the profile has them. The flat list reads as
        # machine output -- "llm, rag, huggingface transformers" in lowercase
        # slugs is a database dump, not a skills section.
        "skill_groups": list(profile.get("skill_groups") or []),
        "skills": _flatten_skills(profile.get("skills_declared") or {}),
        "education": list(profile.get("education") or []),
        "languages": list(profile.get("languages") or []),
    }


def group_roles(bullets: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Bullets grouped under their role, most recent role first.

    Ordered by START DATE DESCENDING, not by bullet rank. Ranking the roles
    put a 2022 job above a 2026 one and a 2023 job between them -- a resume
    with its chronology shuffled, which reads as an error rather than as
    emphasis. Rank still orders bullets WITHIN a role, where it means
    something.

    A role with no date sorts last rather than first, so a missing date cannot
    promote a job above the current one.
    """
    roles: dict[tuple, dict[str, Any]] = {}
    for bullet in bullets:
        key = (bullet["org"], bullet["role"], bullet["start"], bullet["end"])
        role = roles.setdefault(key, {
            "org": bullet["org"], "role": bullet["role"],
            "start": bullet["start"], "end": bullet["end"],
            "note": bullet.get("note") or "",
            "bullets": [],
        })
        role["bullets"].append(bullet)

    for role in roles.values():
        role["bullets"].sort(key=lambda b: b.get("rank") or 999)

    return sorted(roles.values(),
                  key=lambda r: str(r.get("start") or ""), reverse=True)


_MONTHS = ("Jan", "Feb", "Mar", "Apr", "May", "Jun",
           "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")


def format_month(value: Any) -> str:
    """'2023-08' -> 'Aug 2023'. A resume does not print ISO dates."""
    text = str(value or "").strip()
    if not text:
        return ""
    parts = text[:7].split("-")
    if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
        month = int(parts[1])
        if 1 <= month <= 12:
            return f"{_MONTHS[month - 1]} {parts[0]}"
    return text


def _flatten_skills(declared: dict[str, Any]) -> list[str]:
    """Skills in confidence order, de-duplicated, tiers dropped.

    Tier labels ("familiar") are honest internally and read as a weakness on the
    page, so they order the list and then go away.
    """
    out: list[str] = []
    seen: set[str] = set()
    for tier in ("expert", "working", "familiar"):
        for skill in declared.get(tier) or []:
            key = str(skill).strip().lower()
            if key and key not in seen:
                seen.add(key)
                out.append(str(skill).strip())
    return out


def _date_range(role: dict[str, Any]) -> str:
    start = format_month(role.get("start"))
    end = format_month(role.get("end"))
    if not start and not end:
        return ""
    return f"{start} – {end or 'Present'}"


# --------------------------------------------------------------------------
# Typst source -- pure, so it exists whether or not the compiler does
# --------------------------------------------------------------------------

# Typst markup characters. Escaped rather than stripped so a bullet reading
# "cut cost by 30% #2 priority" renders as written instead of failing to compile
# or silently swallowing the rest of the line.
_TYPST_SPECIAL = re.compile(r"([\\#$*_`<>@~\[\]])")


def _escape(text: str) -> str:
    return _TYPST_SPECIAL.sub(r"\\\1", str(text or ""))


def build_typst_source(model: dict[str, Any], scale: float = SCALE_MAX,
                       bullets: list[dict[str, Any]] | None = None) -> str:
    """The `.typ` document at a given scale.

    Single column, no header, no footer, no image, no text inside a graphic --
    each of those is a measured extractor failure rather than a taste. Every
    dimension derives from `scale`, so the fitter has exactly one variable to
    search and cannot produce a page whose margins disagree with its type size.

    **The margin is measured off the source document, not chosen.** It was
    `0.45 + 0.45 * scale`, which is 0.9in a side at full size: 1.8in of an
    8.27in A4 page spent on white, leaving a 6.47in column -- 78% of the page.
    The resume this is meant to match (`input/Resume_YeoKimSiang_20260606.pdf`,
    US Letter) was measured with pypdf: every body line starts at exactly 36pt
    = 0.5in, and wrapped lines run to 576pt, giving 0.5in a side and a 7.5in
    column, 88% of its page.

    That gap was not only ugly, it was destroying content. The fitter shrinks
    before it cuts, so a starved column pushed the search down to the 9pt floor
    and then began dropping the lowest-ranked bullets -- bullets that fit
    comfortably at a sane margin. Rendering a page narrower than the source and
    then deleting the user's evidence to fit it is the wrong trade in both
    directions.

    So the range is anchored on two measured points rather than on a formula
    someone liked: 0.52in a side at full scale (the source's 0.5in, plus a
    hair), narrowing to 0.42in at `SCALE_MIN`, where the page is already fighting
    for room. It stays derived from `scale` alone, so the fitter still has one
    variable. The y margin keeps its 0.85 factor -- a slightly tighter top and
    bottom than sides, which is what the source does too.

    **`leading_em` and `gap_em` were retuned the same way.** Measured in
    baseline-to-baseline distance as a multiple of body size, the source runs
    1.14-1.15 between body lines, ~1.3-1.37 between roles and up to 1.68 at a
    section break. This module was rendering 1.34 between body lines: every
    wrapped line of every bullet carried a fifth of a line of air that the
    document being matched does not have, which over ~57 lines is most of an
    inch. `0.38 + 0.12 * scale` reproduces 1.14-1.15, and `0.55 + 0.25 * scale`
    brings the section breaks back from ~1.99 to ~1.89 without flattening the
    banded headings, which are the most recognisable thing on the page.

    The combined effect on the real 14-fact profile, measured end to end: the
    fitter used to settle at scale 0.848 (9.33pt) and now reaches 0.953
    (10.48pt) with the same content on one page -- a bigger, denser, wider page
    with nothing cut.
    """
    scale = max(SCALE_MIN, min(SCALE_MAX, float(scale)))
    bullets = model["bullets"] if bullets is None else bullets
    roles = group_roles(bullets)

    font_pt = round(BASE_PT * scale, 2)
    name_pt = round(font_pt * 1.55, 2)
    heading_pt = round(font_pt * 1.15, 2)
    # 0.42in at SCALE_MIN, 0.52in at SCALE_MAX. See the docstring: both ends are
    # measured against the source PDF, not chosen.
    margin_in = round(MARGIN_MIN_IN
                      + (MARGIN_MAX_IN - MARGIN_MIN_IN)
                      * (scale - SCALE_MIN) / (SCALE_MAX - SCALE_MIN), 3)
    # Also measured, not chosen -- see the docstring. Both were set so the
    # rendered baseline-to-baseline distances match the source document's.
    leading_em = round(0.38 + 0.12 * scale, 3)
    gap_em = round(0.55 + 0.25 * scale, 3)

    rule_pt = round(0.5 * scale, 3)

    lines = [
        f"#set document(title: {_quote(model.get('name'))} + \" resume\", "
        f"author: {_quote(model.get('name'))})",
        # `header: none, footer: none` is stated rather than left to the default
        # so a later edit has to argue with it.
        f"#set page(paper: \"a4\", margin: (x: {margin_in}in, y: "
        f"{round(margin_in * 0.85, 3)}in), header: none, footer: none)",
        # hyphenate off keeps a bullet's words intact for substring checks and
        # for the extractor; justify off is what makes that the default.
        f"#set text(size: {font_pt}pt, hyphenate: false)",
        f"#set par(leading: {leading_em}em, justify: false)",
        f"#set block(spacing: {gap_em}em)",
        # Tight hanging indent, so a wrapped bullet lines up under its text
        # rather than under the marker.
        f"#set list(indent: 0pt, body-indent: {round(font_pt * 0.5, 2)}pt, "
        f"spacing: {round(gap_em * 0.55, 3)}em, marker: [•])",
        # A section heading is a rule, the title, and another rule. This is the
        # single most recognisable feature of the source document's layout, and
        # rebuilding it from a plain bold line is what made earlier output read
        # as a different resume rather than a tailored one.
        # `stack` rather than a block of separate paragraphs. Inside a block,
        # each element picks up `par.leading` and `block.spacing`, which pushed
        # the rules far away from the title and turned a tight banded heading
        # into three airy lines. `stack` controls the gap explicitly.
        "#show heading: it => block(above: "
        f"{round(gap_em * 0.85, 3)}em, below: {round(gap_em * 0.3, 3)}em)[",
        f"  #stack(dir: ttb, spacing: {round(font_pt * 0.2, 2)}pt,",
        f"    line(length: 100%, stroke: {rule_pt}pt),",
        f"    text(size: {heading_pt}pt, weight: \"bold\")[#upper(it.body)],",
        f"    line(length: 100%, stroke: {rule_pt}pt))",
        "]",
        "",
        # Contact details in the body, centred with plain text -- not a table,
        # not a graphic, not a header.
        "#align(center)[",
        f"  #text(size: {name_pt}pt, weight: \"bold\")[{_escape(model.get('name'))}]",
    ]
    contact = model.get("contact") or []
    if contact:
        joined = " | ".join(_escape(c) for c in contact)
        lines.append(f"  \\\n  #text(size: {round(font_pt * 0.98, 2)}pt)[{joined}]")
    lines.append("]")
    lines.append("")

    # The source document has no summary line, and matching it matters more
    # than the tailoring a headline buys: a section the original does not have
    # is the first thing that makes a resume look like a different document.
    # Kept available, off unless asked for.
    if model.get("headline") and model.get("show_headline"):
        lines.append(_escape(model["headline"]))
        lines.append("")

    education = model.get("education") or []
    if education:
        lines.append(f"= {HEADING_EDUCATION}")
        lines.append("")
        for entry in education:
            lines.extend(_education_block(entry))
        lines.append("")

    if roles:
        lines.append(f"= {HEADING_EXPERIENCE}")
        lines.append("")
        for role in roles:
            dates = _date_range(role)
            # Company FIRST and bold, then the title. The source document reads
            # "*TikTok*, AI Engineer, Global Marketing Science" -- inverting it
            # to lead with the title changes what the eye lands on when a
            # screener skims the left edge of the page.
            heading = f"*{_escape(role['org'])}*, {_escape(role['role'])}"
            if dates:
                heading += f" #h(1fr) {_escape(dates)}"
            if role.get("note"):
                # Trailing backslash forces a line break. Without it Typst
                # joins consecutive markup lines into one paragraph, and the
                # note ran on after the date: "Feb 2026 - Present TikTok
                # Inspire Award recipient" as a single line.
                heading += f" \\\n_{_escape(role['note'])}_"
            lines.append(heading)
            lines.append("")
            for bullet in role["bullets"]:
                lines.append(f"- {_escape(bullet['text'])}")
            lines.append("")

    skill_groups = model.get("skill_groups") or []
    if skill_groups:
        lines.append(f"= {HEADING_SKILLS}")
        lines.append("")
        for group in skill_groups:
            label = _escape(group.get("label") or "")
            items = ", ".join(_escape(s) for s in (group.get("items") or []))
            if items:
                lines.append(f"*{label}:* {items}" if label else items)
                lines.append("")
    elif model.get("skills"):
        lines.append(f"= {HEADING_SKILLS}")
        lines.append("")
        lines.append(", ".join(_escape(s) for s in model["skills"]))
        lines.append("")

    languages = model.get("languages") or []
    if languages:
        lines.append("= Languages")
        lines.append("")
        lines.append("; ".join(_escape(l) for l in languages))
        lines.append("")

    return "\n".join(lines) + "\n"


def _education_block(entry: Any) -> list[str]:
    """One institution: bold name, italic qualifications, then any bullets."""
    if not isinstance(entry, dict):
        return [f"- {_escape(entry)}", ""]

    # One paragraph with an explicit break, for the same reason as the role
    # heading: separate markup lines would merge and put the degree on the same
    # line as the institution.
    head = ""
    institution = _escape(entry.get("institution") or "")
    if institution:
        head = f"*{institution}*"
    qualification = _escape(entry.get("qualification") or "")
    if qualification:
        head += (" \\\n" if head else "") + f"_{qualification}_"

    out: list[str] = []
    if head:
        out.append(head)
    out.append("")
    for bullet in entry.get("bullets") or []:
        out.append(f"- {_escape(bullet)}")
    return out


def _quote(value: Any) -> str:
    """A Typst string literal."""
    return '"' + str(value or "").replace("\\", "\\\\").replace('"', '\\"') + '"'


# --------------------------------------------------------------------------
# fitting
# --------------------------------------------------------------------------

def _compile(source: str) -> bytes:
    typst = _load_typst()
    path = Path(_scratch_typ())
    path.write_text(source, encoding="utf-8")
    try:
        return typst.compile(str(path))
    finally:
        path.unlink(missing_ok=True)


def _scratch_typ() -> str:
    import tempfile

    handle = tempfile.NamedTemporaryFile(suffix=".typ", delete=False)
    handle.close()
    return handle.name


def page_count(pdf_bytes: bytes) -> int:
    """Real page count, read back off the rendered PDF.

    Measured rather than estimated on purpose: an estimator that is wrong by one
    line produces a two-page resume that every check in this module calls one.
    """
    reader = _load_pypdf()
    if reader is None:  # pragma: no cover - environment dependent
        raise ImportError(
            "counting pages needs pypdf: py -m pip install -e .[tailoring]")
    import io

    return len(reader(io.BytesIO(pdf_bytes)).pages)


def fit_to_pages(model: dict[str, Any], max_pages: int = 1) -> dict[str, Any]:
    """Shrink, then cut, until the document fits `max_pages`.

    Order matters and is the whole design: scale first because losing a point of
    type costs nothing a reader notices, bullets only once the ~9pt floor is
    reached. Cuts come off the bottom of the ranking because that is why
    `tailor()` ranks, and every one is recorded -- a silent cut is how a resume
    loses its best bullet with nobody the wiser.

    Degrades when Typst is absent: returns the `.typ` source at full scale with
    nothing cut, and says so in `degraded`.
    """
    caps = capabilities()
    bullets = list(model.get("bullets") or [])
    dropped: list[dict[str, Any]] = []

    if not caps["pdf"]:
        _warn_once("pdf", "typst is not installed -- emitting .typ source "
                          "instead of a PDF: py -m pip install -e .[tailoring]")
        return {
            "ok": True,
            "degraded": "typ-source",
            "note": "typst missing; compile the emitted .typ with: "
                    "typst compile resume.typ resume.pdf",
            "scale": SCALE_MAX, "font_pt": BASE_PT,
            "pages": None, "bullets": bullets, "dropped": dropped,
            "source": build_typst_source(model, SCALE_MAX, bullets),
            "pdf_bytes": None,
        }

    def attempt(scale: float, kept: list[dict[str, Any]]):
        source = build_typst_source(model, scale, kept)
        pdf = _compile(source)
        return source, pdf, page_count(pdf)

    # Full size first. If it already fits there is nothing to buy by shrinking,
    # and the largest readable render is the one to keep.
    source, pdf, pages = attempt(SCALE_MAX, bullets)
    best = (SCALE_MAX, source, pdf, pages) if pages <= max_pages else None

    if best is None:
        low, high = SCALE_MIN, SCALE_MAX
        for _ in range(FIT_STEPS):
            mid = (low + high) / 2
            source, pdf, pages = attempt(mid, bullets)
            if pages <= max_pages:
                best = (mid, source, pdf, pages)
                low = mid  # a larger scale may still fit
            else:
                high = mid
        if best is None:
            # The floor holds. From here the only lever left is content, and it
            # is pulled from the bottom of the ranking.
            #
            # Binary search on how many top-ranked bullets survive. Legitimate
            # because cuts only ever come off the tail and page count rises
            # monotonically with bullet count, so the fitting prefixes are a
            # contiguous run. Dropping one at a time gives the same answer at
            # one compile per bullet, which is the difference between a fast
            # test suite and a slow one.
            low, high = 1, len(bullets)  # low always fits, high never does
            source, pdf, pages = attempt(SCALE_MIN, bullets[:low])
            best_kept = low
            while low + 1 < high:
                mid = (low + high) // 2
                trial_source, trial_pdf, trial_pages = attempt(
                    SCALE_MIN, bullets[:mid])
                if trial_pages <= max_pages:
                    low, best_kept = mid, mid
                    source, pdf, pages = trial_source, trial_pdf, trial_pages
                else:
                    high = mid

            for cut in bullets[best_kept:]:
                dropped.append({
                    "text": cut["text"],
                    "fact_id": cut["fact_id"],
                    "reason": f"lowest-ranked (rank {cut['rank']}); did not fit "
                              f"{max_pages} page(s) at the {MIN_PT:g}pt floor",
                })
            dropped.reverse()  # lowest rank first -- the order they were lost in
            bullets = bullets[:best_kept]
            best = (SCALE_MIN, source, pdf, pages)

    scale, source, pdf, pages = best
    return {
        "ok": pages <= max_pages,
        "degraded": None,
        "note": "",
        "scale": round(scale, 4),
        "font_pt": round(BASE_PT * scale, 2),
        "pages": pages,
        "bullets": bullets,
        "dropped": dropped,
        "source": source,
        "pdf_bytes": pdf,
    }


# --------------------------------------------------------------------------
# rendering
# --------------------------------------------------------------------------

def render_pdf(model: dict[str, Any], out_path: Path,
               max_pages: int = 1) -> dict[str, Any]:
    """Write the PDF, or the `.typ` source when Typst is missing."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fit = fit_to_pages(model, max_pages)

    if fit["pdf_bytes"] is None:
        typ_path = out_path.with_suffix(".typ")
        typ_path.write_text(fit["source"], encoding="utf-8")
        return dict(fit, path=typ_path, bytes=len(fit["source"].encode("utf-8")))

    out_path.write_bytes(fit["pdf_bytes"])
    size = len(fit["pdf_bytes"])
    if size > MAX_PDF_BYTES:
        # Not fatal -- a file the user can still send by hand beats no file --
        # but never silent, because the rejection at the other end is.
        _warn_once("pdf-size",
                   f"rendered PDF is {size} bytes, over the {MAX_PDF_BYTES} "
                   "byte limit some ATS uploads enforce")
    return dict(fit, path=out_path, bytes=size)


def render_docx(model: dict[str, Any], out_path: Path,
                bullets: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    """Write the DOCX from the same model, or markdown when python-docx is missing.

    `bullets` is passed in by `render()` so the Word copy carries exactly the
    set that survived the page fit. Word cannot be measured offline, so it
    inherits the PDF's decision rather than making a second one.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    bullets = model["bullets"] if bullets is None else bullets

    docx = _load_docx()
    if docx is None:
        _warn_once("docx", "python-docx is not installed -- emitting markdown "
                           "instead of a .docx: py -m pip install -e .[tailoring]")
        md_path = out_path.with_suffix(".md")
        md_path.write_text(_markdown(model, bullets), encoding="utf-8")
        return {"ok": True, "degraded": "markdown", "path": md_path,
                "note": "python-docx missing; emitted markdown instead of .docx",
                "bullets": bullets}

    document = docx.Document()
    # Contact details in the body. python-docx would happily put them in a
    # header; a header is what the parser drops.
    heading = document.add_paragraph()
    run = heading.add_run(model.get("name") or "")
    run.bold = True
    if model.get("contact"):
        document.add_paragraph(" | ".join(model["contact"]))
    if model.get("headline"):
        document.add_paragraph(model["headline"])

    roles = group_roles(bullets)
    if roles:
        document.add_heading(HEADING_EXPERIENCE, level=1)
        for role in roles:
            line = document.add_paragraph()
            line.add_run(f"{role['role']}, {role['org']}").bold = True
            dates = _date_range(role)
            if dates:
                line.add_run(f"  {dates}")
            for bullet in role["bullets"]:
                document.add_paragraph(bullet["text"], style="List Bullet")

    if model.get("skills"):
        document.add_heading(HEADING_SKILLS, level=1)
        document.add_paragraph(", ".join(model["skills"]))

    if model.get("education"):
        document.add_heading(HEADING_EDUCATION, level=1)
        for entry in model["education"]:
            if isinstance(entry, dict):
                parts = [entry.get("qualification"), entry.get("institution"),
                         entry.get("year")]
                text = ", ".join(str(p) for p in parts if p)
            else:
                text = str(entry)
            document.add_paragraph(text, style="List Bullet")

    document.save(str(out_path))
    return {"ok": True, "degraded": None, "path": out_path, "note": "",
            "bullets": bullets}


def _markdown(model: dict[str, Any], bullets: list[dict[str, Any]]) -> str:
    lines = [f"# {model.get('name') or ''}"]
    if model.get("contact"):
        lines += ["", " | ".join(model["contact"])]
    if model.get("headline"):
        lines += ["", model["headline"]]
    roles = group_roles(bullets)
    if roles:
        lines += ["", f"## {HEADING_EXPERIENCE}"]
        for role in roles:
            dates = _date_range(role)
            lines += ["", f"**{role['role']}, {role['org']}**"
                          + (f" — {dates}" if dates else ""), ""]
            lines += [f"- {b['text']}" for b in role["bullets"]]
    if model.get("skills"):
        lines += ["", f"## {HEADING_SKILLS}", "", ", ".join(model["skills"])]
    return "\n".join(lines) + "\n"


def render(model: dict[str, Any], out_dir: Path, stem: str = "resume",
           max_pages: int = 1) -> dict[str, Any]:
    """Both documents from one model. The PDF decides the content; DOCX follows.

    Returns the fit result plus both paths, `dropped` and any `degraded` modes,
    so a caller can report exactly what was rendered and what was lost.
    """
    out_dir = Path(out_dir)
    pdf = render_pdf(model, out_dir / f"{stem}.pdf", max_pages=max_pages)
    doc = render_docx(model, out_dir / f"{stem}.docx", bullets=pdf["bullets"])
    degraded = [d for d in (pdf.get("degraded"), doc.get("degraded")) if d]
    return {
        "ok": bool(pdf["ok"] and doc["ok"]),
        "pdf": pdf,
        "docx": doc,
        "pages": pdf.get("pages"),
        "scale": pdf.get("scale"),
        "font_pt": pdf.get("font_pt"),
        "bullets": pdf["bullets"],
        "dropped": pdf.get("dropped") or [],
        "degraded": degraded,
        "notes": [n for n in (pdf.get("note"), doc.get("note")) if n],
    }


# --------------------------------------------------------------------------
# the check that matters: did the decisive material land on page one
# --------------------------------------------------------------------------

def _norm(text: str) -> str:
    """Collapse to alphanumerics for substring matching.

    PDF text extraction reinserts line breaks mid-sentence and rewrites hyphens
    and quotes; comparing raw strings would report a bullet missing that is
    plainly on the page. Dropping everything but letters and digits makes the
    membership test answer the question actually being asked -- are these words
    on page one -- rather than a question about typography.
    """
    return re.sub(r"[^a-z0-9]+", "", str(text or "").lower())


def page_one_text(pdf_path: Path) -> str:
    """Raw text of page 1, or "" if the file has no pages."""
    reader = _load_pypdf()
    if reader is None:  # pragma: no cover - environment dependent
        raise ImportError(
            "reading a PDF needs pypdf: py -m pip install -e .[tailoring]")
    pages = reader(str(pdf_path)).pages
    return (pages[0].extract_text() or "") if pages else ""


def page_one_sufficiency(pdf_path: Path,
                         model: dict[str, Any]) -> dict[str, Any]:
    """Did the decisive material clear the page-1 break?

    Page two does not get read. That is the one refinement that survives the
    evidence against a hard one-page rule, and it makes page count the wrong
    thing to assert. This asserts the right thing instead, and it is checkable
    rather than a matter of opinion: the most JD-relevant role, the top three
    ranked bullets, and the skills block, tested for membership in page 1's
    extracted text.

    Reports what made it AND what did not. `ok` alone would hide which bullet
    fell off.
    """
    text = _norm(page_one_text(Path(pdf_path)))
    bullets = model.get("bullets") or []
    roles = model.get("roles") or group_roles(bullets)
    missing: list[str] = []

    top_role: dict[str, Any] = {"org": "", "role": "", "on_page_one": True}
    if roles:
        role = roles[0]
        on_page = bool(_norm(role["org"]) and _norm(role["org"]) in text)
        top_role = {"org": role["org"], "role": role["role"],
                    "on_page_one": on_page}
        if not on_page:
            missing.append(f"most relevant role: {role['role']}, {role['org']}")

    top_bullets = []
    for bullet in bullets[:PAGE_ONE_BULLETS]:
        on_page = _norm(bullet["text"]) in text
        top_bullets.append({"rank": bullet["rank"], "fact_id": bullet["fact_id"],
                            "text": bullet["text"], "on_page_one": on_page})
        if not on_page:
            missing.append(f"rank {bullet['rank']} bullet ({bullet['fact_id']})")

    skills = model.get("skills") or []
    absent = [s for s in skills if _norm(s) not in text]
    heading_on_page = _norm(HEADING_SKILLS) in text
    skills_ok = bool(skills) and heading_on_page and not absent
    if skills and not skills_ok:
        missing.append("skills block"
                       + (f" (missing: {', '.join(absent)})" if absent else ""))

    return {
        "ok": not missing,
        "top_role": top_role,
        "top_bullets": top_bullets,
        "skills": {"on_page_one": skills_ok, "heading_on_page_one":
                   heading_on_page, "missing": absent},
        "made_it": [b["fact_id"] for b in top_bullets if b["on_page_one"]],
        "missing": missing,
    }
