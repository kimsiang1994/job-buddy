"""Tests for rendering a tailored model into a PDF and a DOCX.

    py -m unittest tests.test_render_resume

Offline, no API key, no network -- the fixtures are a hand-written model, not a
live `tailor()` call, so nothing here depends on what a model happened to
choose today.

Three properties are worth the render cost and none of them can be read off the
source. Whether a page fits is a fact about the compiled PDF, so it is asserted
by re-reading the PDF with pypdf rather than by trusting the fitter. Whether a
cut took the *right* bullet is the difference between a resume that lost its
weakest line and one that lost its best, and only the ranking distinguishes
them. And whether the decisive material cleared the page-1 break is the check
that survives the evidence against a hard one-page rule -- page two does not get
read, so page count is the wrong assertion and membership in page 1 is the right
one.

Every person, employer and achievement below is invented. This repo is public
and must never carry real CV data.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from jobbuddy import render_resume

IDENTITY = {
    "name": "Ada Nakamura",
    "email": "ada.nakamura@example.com",
    "phone": "+65 8000 0000",
    "location": "Singapore",
    "links": ["github.com/example-ada"],
}

SKILLS = {"expert": ["Python", "PySpark"], "working": ["Airflow"],
          "familiar": ["Kubernetes"]}

PROFILE = {"identity": IDENTITY, "skills_declared": SKILLS}


def bullet(text: str, fact_id: str, org: str = "Umbra Financial",
           role: str = "Manager, Data Engineering",
           start: str = "2023-08", end: str = "2024-11") -> dict:
    return {"text": text, "fact_id": fact_id, "org": org, "role": role,
            "start": start, "end": end}


# A plausible tailored draft: short, ranked, comfortably one page.
TAILORED = {
    "headline": "Data engineer building measured retrieval and ETL systems.",
    "bullets": [
        bullet("Built a retrieval pipeline serving 12 million daily requests in "
               "PyTorch, cutting p99 latency from 340ms to 90ms",
               "northwind.retrieval", "Northwind Labs", "AI Engineer",
               "2024-08", ""),
        bullet("Cut data generation from 10 to 5 working days by automating 4 "
               "ETL processes in PySpark", "umbra.etl"),
        bullet("Rebuilt the nightly reconciliation job to run in 20 minutes",
               "umbra.reconciliation"),
        bullet("Introduced schema contracts across 6 upstream feeds",
               "umbra.contracts"),
    ],
}


def overflowing(count: int = 40) -> dict:
    """A draft too long for one page even at the type-size floor."""
    bullets = [TAILORED["bullets"][0]]
    bullets += [
        bullet(f"Filler achievement {i} " + " ".join(
            ["describing a measured outcome across the platform"] * 4),
            f"umbra.filler.{i}")
        for i in range(2, count + 1)
    ]
    return {"headline": TAILORED["headline"], "bullets": bullets}


def split_across_pages() -> dict:
    """Ranks 2 and 3 sit in the OLDER role, behind a very long recent one.

    A real layout outcome, not a contrivance: roles render newest first and
    bullets are grouped under them, so a high-ranked bullet belonging to an
    earlier job is pushed down by every bullet of the current one. Exactly the
    case where page count says "fine" and the screener never sees the
    candidate's second-best evidence.

    The filler deliberately sits in the RECENT role. An earlier version of this
    fixture put it in the older one and relied on roles being ordered by bullet
    rank -- which was itself the chronology bug, so the fixture only reproduced
    the failure while that bug was present.
    """
    recent = ("Northwind Labs", "AI Engineer", "2024-08", "")
    bullets = [
        bullet("Built a retrieval pipeline serving 12 million daily requests "
               "in PyTorch", "northwind.retrieval", *recent),
    ]
    bullets += [
        bullet("Cut data generation from 10 to 5 working days by "
               "automating 4 ETL processes in PySpark", "umbra.etl"),
        bullet("Built data-quality checks that halved debugging time",
               "umbra.quality"),
    ]
    bullets += [
        bullet(f"Filler achievement {i} " + " ".join(
            ["describing a measured outcome across the platform"] * 4),
            f"northwind.filler.{i}", *recent)
        for i in range(4, 40)
    ]
    return {"headline": TAILORED["headline"], "bullets": bullets}


def model_for(tailored: dict) -> dict:
    return render_resume.build_model(PROFILE, tailored)


def flatten(text: str) -> str:
    """Alphanumerics only -- see `render_resume._norm`."""
    return "".join(c for c in text.lower() if c.isalnum())


class RenderCase(unittest.TestCase):
    """Gives each test a scratch directory that cleans itself up."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.out = Path(self._tmp.name)
        self.addCleanup(self._tmp.cleanup)
        self.addCleanup(render_resume.reset_capability_cache)
        render_resume.reset_capability_cache()
        if not render_resume.capabilities()["pdf"]:
            self.skipTest("typst is not installed; PDF rendering untestable")


class TheModelIsTheSingleSourceOfTruth(unittest.TestCase):
    def test_bullets_keep_the_rank_order_tailor_returned(self):
        model = model_for(TAILORED)
        self.assertEqual([b["rank"] for b in model["bullets"]], [1, 2, 3, 4])
        self.assertEqual(model["bullets"][0]["fact_id"], "northwind.retrieval")

    def test_roles_are_reverse_chronological_not_rank_ordered(self):
        """Ranking roles produced a resume running 2022, 2026, 2023 -- a
        shuffled work history, which a reader takes for an error rather than
        for emphasis. Rank orders bullets WITHIN a role, where it means
        something; date orders the roles."""
        model = model_for(TAILORED)
        self.assertEqual([r["org"] for r in model["roles"]],
                         ["Northwind Labs", "Umbra Financial"])

    def test_the_most_recent_role_leads_even_when_it_ranks_lower(self):
        """The case the previous test cannot distinguish, because there the
        top-ranked role is also the most recent."""
        tailored = {"headline": "", "bullets": [
            bullet("Cut data generation from 10 to 5 working days",
                   "umbra.etl"),
            bullet("Built a retrieval pipeline serving 12 million requests",
                   "northwind.retrieval", "Northwind Labs", "AI Engineer",
                   "2024-08", ""),
        ]}
        model = model_for(tailored)
        self.assertEqual(model["roles"][0]["org"], "Northwind Labs")

    def test_every_employer_survives_into_the_model(self):
        """A job dropped from the resume is a gap in the work history."""
        model = model_for(TAILORED)
        self.assertEqual({r["org"] for r in model["roles"]},
                         {"Northwind Labs", "Umbra Financial"})

    def test_contact_details_live_in_the_model_body_not_a_header(self):
        model = model_for(TAILORED)
        self.assertIn("ada.nakamura@example.com", model["contact"])
        self.assertIn("Singapore", model["contact"])

    def test_skill_tier_labels_do_not_reach_the_page(self):
        model = model_for(TAILORED)
        self.assertEqual(model["skills"],
                         ["Python", "PySpark", "Airflow", "Kubernetes"])


class ThePageBudgetIsEnforcedAgainstTheRealPDF(RenderCase):
    def test_the_default_budget_renders_exactly_one_page(self):
        result = render_resume.render(model_for(TAILORED), self.out)
        self.assertTrue(result["ok"])
        self.assertEqual(
            render_resume.page_count(result["pdf"]["path"].read_bytes()), 1)

    def test_a_draft_that_overflows_is_still_forced_onto_one_page(self):
        result = render_resume.render(model_for(overflowing()), self.out)
        self.assertEqual(
            render_resume.page_count(result["pdf"]["path"].read_bytes()), 1)

    def test_max_pages_two_permits_a_second_page(self):
        """The budget is a parameter. A hard one-page rule is folklore at 5+ years."""
        result = render_resume.render(model_for(overflowing()), self.out,
                                      max_pages=2)
        pages = render_resume.page_count(result["pdf"]["path"].read_bytes())
        self.assertEqual(pages, 2)
        self.assertTrue(result["ok"])

    def test_a_larger_budget_keeps_bullets_a_one_page_budget_would_cut(self):
        """Otherwise `max_pages` would be decorative."""
        one = render_resume.render(model_for(overflowing()), self.out, "one", 1)
        two = render_resume.render(model_for(overflowing()), self.out, "two", 2)
        self.assertGreater(len(two["bullets"]), len(one["bullets"]))

    def test_shrinking_is_tried_before_anything_is_cut(self):
        """A point of type size costs a reader nothing; a lost bullet costs them."""
        result = render_resume.render(model_for(overflowing(20)), self.out)
        self.assertLess(result["font_pt"], render_resume.BASE_PT)
        self.assertEqual(result["dropped"], [])


class TheTypeSizeFloorHolds(RenderCase):
    def test_no_render_goes_below_the_floor(self):
        for count in (4, 20, 40):
            with self.subTest(bullets=count):
                result = render_resume.render(model_for(overflowing(count)),
                                              self.out, f"s{count}")
                self.assertGreaterEqual(result["font_pt"], render_resume.MIN_PT)

    def test_the_floor_is_what_forces_the_cut(self):
        result = render_resume.render(model_for(overflowing()), self.out)
        self.assertAlmostEqual(result["font_pt"], render_resume.MIN_PT, places=2)
        self.assertTrue(result["dropped"])


class CutsComeOffTheBottomOfTheRankingAndAreRecorded(RenderCase):
    def test_the_lowest_ranked_bullets_go_first(self):
        tailored = overflowing()
        result = render_resume.render(model_for(tailored), self.out)
        kept = len(result["bullets"])
        expected = [b["fact_id"] for b in tailored["bullets"][kept:]]
        self.assertEqual([d["fact_id"] for d in result["dropped"]],
                         list(reversed(expected)))

    def test_the_top_ranked_bullet_survives_every_cut(self):
        result = render_resume.render(model_for(overflowing()), self.out)
        self.assertEqual(result["bullets"][0]["fact_id"], "northwind.retrieval")
        self.assertNotIn("northwind.retrieval",
                         [d["fact_id"] for d in result["dropped"]])

    def test_every_cut_carries_its_text_fact_id_and_reason(self):
        """A silent cut is how a resume loses a bullet with nobody the wiser."""
        result = render_resume.render(model_for(overflowing()), self.out)
        self.assertTrue(result["dropped"])
        for record in result["dropped"]:
            with self.subTest(fact_id=record["fact_id"]):
                self.assertTrue(record["text"])
                self.assertTrue(record["fact_id"])
                self.assertIn("lowest-ranked", record["reason"])

    def test_nothing_is_cut_when_the_draft_already_fits(self):
        result = render_resume.render(model_for(TAILORED), self.out)
        self.assertEqual(result["dropped"], [])
        self.assertEqual(len(result["bullets"]), 4)


class PageOneIsWhatActuallyGetsRead(RenderCase):
    def test_a_fitting_resume_lands_everything_decisive_on_page_one(self):
        model = model_for(TAILORED)
        result = render_resume.render(model, self.out)
        report = render_resume.page_one_sufficiency(result["pdf"]["path"], model)
        self.assertTrue(report["ok"], report["missing"])
        self.assertTrue(report["top_role"]["on_page_one"])
        self.assertTrue(all(b["on_page_one"] for b in report["top_bullets"]))
        self.assertTrue(report["skills"]["on_page_one"])

    def test_a_top_bullet_pushed_below_the_break_is_reported_by_rank(self):
        """Page count says two pages is fine. It does not say the rank 2 bullet
        is on the page nobody reads."""
        model = model_for(split_across_pages())
        result = render_resume.render(model, self.out, max_pages=2)
        self.assertEqual(result["pages"], 2)

        report = render_resume.page_one_sufficiency(result["pdf"]["path"], model)
        self.assertFalse(report["ok"])
        by_rank = {b["rank"]: b for b in report["top_bullets"]}
        self.assertTrue(by_rank[1]["on_page_one"])
        self.assertFalse(by_rank[2]["on_page_one"])
        self.assertIn("rank 2 bullet (umbra.etl)", report["missing"])

    def test_the_report_says_what_made_it_as_well_as_what_did_not(self):
        model = model_for(split_across_pages())
        result = render_resume.render(model, self.out, max_pages=2)
        report = render_resume.page_one_sufficiency(result["pdf"]["path"], model)
        self.assertEqual(report["made_it"], ["northwind.retrieval"])
        self.assertTrue(report["missing"])

    def test_a_skills_block_below_the_break_is_reported(self):
        model = model_for(split_across_pages())
        result = render_resume.render(model, self.out, max_pages=2)
        report = render_resume.page_one_sufficiency(result["pdf"]["path"], model)
        self.assertFalse(report["skills"]["on_page_one"])
        self.assertIn("Python", report["skills"]["missing"])


class TheFileSurvivesAnATSUpload(RenderCase):
    def test_the_pdf_stays_under_the_100kb_limit(self):
        """Taleo documents a 100 KB cap; over it the upload is rejected."""
        for name, tailored in (("short", TAILORED), ("long", overflowing())):
            with self.subTest(draft=name):
                result = render_resume.render(model_for(tailored), self.out, name)
                size = result["pdf"]["path"].stat().st_size
                self.assertLess(size, render_resume.MAX_PDF_BYTES,
                                f"{name} rendered {size} bytes")

    def test_conventional_section_headings_reach_the_page(self):
        """Unconventional headings are the documented per-vendor parse failure."""
        model = model_for(TAILORED)
        result = render_resume.render(model, self.out)
        text = flatten(render_resume.page_one_text(result["pdf"]["path"]))
        self.assertIn(flatten("Experience"), text)
        self.assertIn(flatten("Skills"), text)

    def test_contact_details_are_extractable_from_the_body(self):
        model = model_for(TAILORED)
        result = render_resume.render(model, self.out)
        text = flatten(render_resume.page_one_text(result["pdf"]["path"]))
        self.assertIn(flatten("ada.nakamura@example.com"), text)


class BothDocumentsSayTheSameThing(RenderCase):
    def test_pdf_and_docx_carry_identical_bullet_text(self):
        """The failure this prevents: someone edits one template and ships a
        Word copy three bullets behind the PDF."""
        docx = self._require_docx()
        model = model_for(TAILORED)
        result = render_resume.render(model, self.out)

        pdf_text = flatten(render_resume.page_one_text(result["pdf"]["path"]))
        paragraphs = [p.text for p in
                      docx.Document(str(result["docx"]["path"])).paragraphs]
        docx_text = flatten(" ".join(paragraphs))

        for line in result["bullets"]:
            with self.subTest(fact_id=line["fact_id"]):
                self.assertIn(flatten(line["text"]), pdf_text)
                self.assertIn(flatten(line["text"]), docx_text)

    def test_a_bullet_cut_from_the_pdf_is_cut_from_the_docx_too(self):
        """Word cannot be measured offline, so it must inherit the PDF's cut
        rather than make a second, different one."""
        docx = self._require_docx()
        result = render_resume.render(model_for(overflowing()), self.out)
        self.assertTrue(result["dropped"])

        paragraphs = [p.text for p in
                      docx.Document(str(result["docx"]["path"])).paragraphs]
        docx_text = flatten(" ".join(paragraphs))
        for record in result["dropped"]:
            with self.subTest(fact_id=record["fact_id"]):
                self.assertNotIn(flatten(record["text"]), docx_text)

    def _require_docx(self):
        if not render_resume.capabilities()["docx"]:
            self.skipTest("python-docx is not installed")
        import docx

        return docx


class AMissingWheelDegradesAndNeverRaises(unittest.TestCase):
    """Core is stdlib. An optional dependency is optional or it is not."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.out = Path(self._tmp.name)
        self.addCleanup(self._tmp.cleanup)
        self.addCleanup(render_resume.reset_capability_cache)
        self.addCleanup(setattr, render_resume, "_load_typst",
                        render_resume._load_typst)
        self.addCleanup(setattr, render_resume, "_load_docx",
                        render_resume._load_docx)

    def _without(self, name: str) -> None:
        setattr(render_resume, f"_load_{name}", lambda: None)
        render_resume.reset_capability_cache()

    def test_no_typst_emits_typ_source_and_names_the_compile_command(self):
        self._without("typst")
        with self.assertWarns(RuntimeWarning):
            result = render_resume.render_pdf(model_for(TAILORED),
                                              self.out / "resume.pdf")
        self.assertEqual(result["degraded"], "typ-source")
        self.assertTrue(result["path"].name.endswith(".typ"))
        self.assertIn("typst compile", result["note"])
        self.assertIn("Experience", result["path"].read_text(encoding="utf-8"))

    def test_no_python_docx_emits_markdown_carrying_the_same_bullets(self):
        self._without("docx")
        with self.assertWarns(RuntimeWarning):
            result = render_resume.render_docx(model_for(TAILORED),
                                               self.out / "resume.docx")
        self.assertEqual(result["degraded"], "markdown")
        text = result["path"].read_text(encoding="utf-8")
        for line in TAILORED["bullets"]:
            self.assertIn(line["text"], text)

    def test_render_reports_which_degradation_was_taken(self):
        self._without("typst")
        self._without("docx")
        with self.assertWarns(RuntimeWarning):
            result = render_resume.render(model_for(TAILORED), self.out)
        self.assertEqual(sorted(result["degraded"]), ["markdown", "typ-source"])
        self.assertTrue(result["notes"])

    def test_the_warning_is_emitted_once_not_once_per_render(self):
        """A per-job warning would bury the line saying which mode was taken."""
        self._without("typst")
        with self.assertWarns(RuntimeWarning):
            render_resume.render_pdf(model_for(TAILORED), self.out / "a.pdf")
        import warnings

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            render_resume.render_pdf(model_for(TAILORED), self.out / "b.pdf")
        self.assertEqual([w for w in caught
                          if issubclass(w.category, RuntimeWarning)], [])

    def test_building_typst_source_never_needs_the_compiler(self):
        self._without("typst")
        source = render_resume.build_typst_source(model_for(TAILORED))
        self.assertIn("Umbra Financial", source)


class TypstMarkupInABulletCannotBreakTheDocument(RenderCase):
    def test_hashes_and_asterisks_survive_into_the_rendered_page(self):
        """A bullet reading '#2 priority, cut cost 30%' must not be markup."""
        tailored = {"headline": "", "bullets": [
            bullet("Ranked #2 priority: cut cost by 30% using C++ and *ETL* "
                   "with fixed_width parsing", "umbra.markup")]}
        model = model_for(tailored)
        result = render_resume.render(model, self.out)
        text = flatten(render_resume.page_one_text(result["pdf"]["path"]))
        self.assertIn(flatten("Ranked #2 priority"), text)
        self.assertIn(flatten("fixed_width"), text)


if __name__ == "__main__":
    unittest.main(verbosity=2)
